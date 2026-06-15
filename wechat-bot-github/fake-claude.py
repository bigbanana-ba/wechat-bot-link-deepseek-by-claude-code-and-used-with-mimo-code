#!/usr/bin/env python3
"""
AI Agent for cc-connect - uses DeepSeek API with tool calling.
Gives the bot the ability to read/write files, run commands, and more.
"""

import sys
import json
import uuid
import urllib.request
import urllib.error
import urllib.parse
import os
import subprocess
import glob
import time
import threading
import re
import html as html_mod
import base64
import http.client
import shutil as _shutil

# MCP client support
try:
    from mcp_client import init_mcp, call_mcp_tool, stop_all_mcp, MCP_SERVERS
    _mcp_available = True
except ImportError:
    _mcp_available = False
    MCP_SERVERS = []

# ============================================================
# Configuration
# ============================================================
DEEPSEEK_API_KEY = "你的DeepSeek-API-Key"
DEEPSEEK_URL = "https://api.deepseek.com/chat/completions"
DEFAULT_MODEL = "deepseek-chat"  # 日常聊天用便宜模型，复杂任务由 AI 调 v4-pro

# Qwen-VL (阿里千问) 视觉识别配置 - 用于图片识别
QWEN_API_KEY = "你的千问-API-Key"
QWEN_BASE_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1"
QWEN_VL_MODEL = "qwen3-vl-plus"

# Attachment monitoring
ATTACH_DIR = os.path.join(os.environ.get("USERPROFILE", "C:\\Users\\big"), ".cc-connect", "attachments")

# Debug log
DEBUG_LOG = os.path.join(os.path.dirname(os.path.abspath(__file__)), "debug.jsonl")

# Safety mode: True = restricted access, False = full access to everything
# Set to False to give the bot unrestricted access to your entire computer
SAFE_MODE = False

# Multi-model configuration
# The bot can switch between models based on task difficulty
MODELS = {
    "deepseek-chat": {
        "name": "DeepSeek V3",
        "api_key": "你的DeepSeek-API-Key",
        "base_url": "https://api.deepseek.com/chat/completions",
        "description": "通用对话，日常任务",
        "level": "normal",
    },
    "deepseek-reasoner": {
        "name": "DeepSeek R1",
        "api_key": "你的DeepSeek-API-Key",
        "base_url": "https://api.deepseek.com/chat/completions",
        "description": "深度推理，数学/逻辑/编程难题",
        "level": "advanced",
    },
    "deepseek-v4-pro": {
        "name": "DeepSeek V4 Pro",
        "api_key": "你的DeepSeek-API-Key",
        "base_url": "https://api.deepseek.com/chat/completions",
        "description": "最强模型，复杂分析/创作/决策",
        "level": "pro",
    },
    "deepseek-v4-flash": {
        "name": "DeepSeek V4 Flash",
        "api_key": "你的DeepSeek-API-Key",
        "base_url": "https://api.deepseek.com/chat/completions",
        "description": "快速响应，简单任务",
        "level": "fast",
    },
}

# Memory storage
MEMORY_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "memory.json")
MAX_API_HISTORY = 400  # 最大历史消息数
SHORT_HISTORY = 40     # 闲聊场景：只需少量上下文
NORMAL_HISTORY = 100   # 正常场景
LONG_HISTORY = 200     # 复杂任务场景
MAX_DISK_HISTORY = 0  # Disk: 0 = unlimited, keep all history forever

# ============================================================
# Memory management
# ============================================================

def load_memory():
    """Load ALL conversation history and user facts from disk."""
    if not os.path.exists(MEMORY_FILE):
        return [], {}

    try:
        with open(MEMORY_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        history = data.get("history", [])
        facts = data.get("facts", {})
        # Validate format
        valid_history = []
        for m in history:
            if isinstance(m, dict) and "role" in m and "content" in m:
                valid_history.append(m)
        sys.stderr.write(f"Memory loaded: {len(valid_history)} messages, {len(facts)} facts\n")
        return valid_history, facts
    except (json.JSONDecodeError, IOError) as e:
        sys.stderr.write(f"Memory load failed: {e}\n")
        return [], {}


def save_memory(history, facts):
    """Save ALL conversation history and facts to disk (thread-safe)."""
    def _save():
        try:
            data = {
                "history": history,  # Keep everything!
                "facts": facts,
                "total_messages": len(history),
                "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            }
            os.makedirs(os.path.dirname(MEMORY_FILE), exist_ok=True)
            with open(MEMORY_FILE, 'w', encoding='utf-8', newline='\n') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
                f.flush()
                os.fsync(f.fileno())
        except IOError as e:
            sys.stderr.write(f"Memory save failed: {e}\n")

    # Fire and forget (non-blocking)
    t = threading.Thread(target=_save, daemon=True)
    t.start()

SYSTEM_PROMPT = """你是 DeepSeek AI 助手。

## 用户偏好
{preferences_text}

## 规则
- 直接回答，不说动作描述
- 简洁，200字内，中文
- 可用Unicode表情😊👍🎉，禁止[旺柴]
- 日常聊天自己答，复杂任务用 ask_model
"""

# Memory facts storage
PREF_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "preferences.txt")
memory_facts = {}

def get_system_prompt():
    """Build system prompt with current preferences."""
    prefs = ""
    # Read preferences from text file
    if os.path.exists(PREF_FILE):
        try:
            with open(PREF_FILE, 'r', encoding='utf-8') as f:
                prefs = f.read().strip()
        except:
            pass
    # Also merge in memory_facts dict
    if memory_facts:
        for k, v in memory_facts.items():
            prefs += f"\n- {k}: {v}"
    if not prefs:
        prefs = "(暂无偏好设置，你可以使用 remember_preference 工具记录用户偏好)"
    return SYSTEM_PROMPT.replace("{preferences_text}", prefs)

# Allowed directories for file operations (safety)
ALLOWED_DIRS = [
    "C:\\Users\\big",
    "C:\\Users\\big\\Desktop",
    "C:\\Users\\big\\Documents",
    "C:\\Users\\big\\Downloads",
    "C:\\Users\\big\\wechat-bot",
]

# Conversation history
conversation = []

# ============================================================
# Tool definitions (OpenAI/DeepSeek function calling format)
# ============================================================
TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "remember_preference",
            "description": "记住用户的偏好或要求。如回复风格、行为规则等。用户说'以后不要xxx'或'回复简洁点'时用这个。",
            "parameters": {
                "type": "object",
                "properties": {
                    "rule": {
                        "type": "string",
                        "description": "要记住的偏好规则，如 '回复不要描述动作'"
                    }
                },
                "required": ["rule"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "forget_preferences",
            "description": "清除部分或全部用户偏好。",
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "要清除的关键词，留空清除全部"
                    }
                },
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "ask_model",
            "description": "将问题交给更强的模型处理。日常聊天不要用，只在写代码、复杂推理、深度分析时用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "model": {"type": "string", "description": "deepseek-v4-pro(最强1M上下文) 或 deepseek-reasoner(推理)"},
                    "question": {"type": "string", "description": "问题"}
                },
                "required": ["model", "question"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "send_to_wechat",
            "description": "发送文件或图片给微信用户。生成文件（如Excel、代码、截图）后用此工具发回微信。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "文件路径"},
                    "type": {"type": "string", "description": "image 或 file", "enum": ["image", "file"]}
                },
                "required": ["path", "type"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_time",
            "description": "获取当前日期、时间、星期。用户问时间/日期/星期几时使用。",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "github_search",
            "description": "搜索 GitHub 仓库和代码。用于找开源项目、看trending、查技术方案。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "搜索关键词"},
                    "type": {"type": "string", "description": "repositories 或 code", "enum": ["repositories", "code"]}
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "搜索网页，获取最新信息。使用 Chrome 浏览器，国内外网站都能访问。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "搜索关键词"}
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "bash",
            "description": "执行 Shell 命令（120秒超时）。可执行 Python 脚本、操作文件、调用系统工具等。",
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {
                        "type": "string",
                        "description": "要执行的命令，如 dir、type、echo 等"
                    }
                },
                "required": ["command"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read",
            "description": "读取一个文件的内容",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "文件的绝对路径"
                    },
                    "limit": {
                        "type": "integer",
                        "description": "最多读取行数，默认200"
                    }
                },
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "write",
            "description": "创建或覆盖一个文件",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "文件的绝对路径"
                    },
                    "content": {
                        "type": "string",
                        "description": "要写入的内容"
                    }
                },
                "required": ["path", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_dir",
            "description": "列出目录中的文件和子目录",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "目录的绝对路径"
                    }
                },
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "grep",
            "description": "在文件中搜索匹配的内容",
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {
                        "type": "string",
                        "description": "要搜索的文本或正则表达式"
                    },
                    "path": {
                        "type": "string",
                        "description": "要搜索的文件或目录路径"
                    }
                },
                "required": ["pattern", "path"]
            }
        }
    },
]


# ============================================================
# Tool implementations
# ============================================================

def is_path_allowed(path):
    """Check if a path is within allowed directories."""
    if not SAFE_MODE:
        return True  # Full access
    abs_path = os.path.abspath(path)
    for allowed in ALLOWED_DIRS:
        if abs_path.lower().startswith(allowed.lower()):
            return True
    return False


def tool_read(path, limit=200):
    """Read a file."""
    if not is_path_allowed(path):
        return f"[拒绝] 无权访问此路径: {path}"
    try:
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            lines = f.readlines()
            if len(lines) > limit:
                return ''.join(lines[:limit]) + f"\n... (共 {len(lines)} 行，仅显示前 {limit} 行)"
            return ''.join(lines) if lines else "(空文件)"
    except FileNotFoundError:
        return f"[错误] 文件不存在: {path}"
    except PermissionError:
        return f"[错误] 无权限读取: {path}"
    except Exception as e:
        return f"[错误] {str(e)}"


def tool_write(path, content):
    """Write a file."""
    if not is_path_allowed(path):
        return f"[拒绝] 无权在此位置创建文件: {path}"
    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return f"文件已写入: {path} ({len(content)} 字符)"
    except PermissionError:
        return f"[错误] 无权限写入: {path}"
    except Exception as e:
        return f"[错误] {str(e)}"


def tool_send_to_wechat(path, ftype="file"):
    """Send a file or image back to WeChat via cc-connect."""
    if not os.path.exists(path):
        return f"[错误] 文件不存在: {path}"
    flag = "--image" if ftype == "image" else "--file"
    # Ensure full path to cc-connect
    cc_path = _shutil.which("cc-connect") or "cc-connect"
    cmd = f'{cc_path} send {flag} "{os.path.abspath(path)}" -p wechat-bot'
    sys.stderr.write(f"[SEND] Running: {cmd}\n")
    sys.stderr.flush()
    try:
        result = subprocess.run(
            cmd,
            shell=True, capture_output=True, text=True, timeout=30,
            encoding='utf-8', errors='replace'
        )
        sys.stderr.write(f"[SEND] rc={result.returncode} out={result.stdout[:100]} err={result.stderr[:100]}\n")
        sys.stderr.flush()
        if result.returncode == 0:
            return f"已发送{ftype}: {os.path.basename(path)}"
        return f"[发送失败] {result.stderr[:200] or result.stdout[:200]}"
    except Exception as e:
        sys.stderr.write(f"[SEND] Exception: {e}\n")
        sys.stderr.flush()
        return f"[发送失败] {str(e)[:100]}"


def tool_remember_preference(rule):
    """Remember a user preference/rule."""
    global memory_facts, conversation
    # Write to human-readable preferences.txt
    try:
        existing = ""
        if os.path.exists(PREF_FILE):
            with open(PREF_FILE, 'r', encoding='utf-8') as f:
                existing = f.read().strip()
        if existing:
            existing += "\n"
        with open(PREF_FILE, 'w', encoding='utf-8') as f:
            f.write(existing + rule)
        memory_facts[rule[:30]] = rule
        save_memory(conversation, memory_facts)
        return f"已记住偏好: {rule}"
    except Exception as e:
        return f"[错误] {str(e)}"


def tool_forget_preferences(keyword=None):
    """Forget preferences."""
    global memory_facts, conversation
    if keyword:
        # Remove matching lines from preferences.txt
        if os.path.exists(PREF_FILE):
            try:
                with open(PREF_FILE, 'r', encoding='utf-8') as f:
                    lines = f.readlines()
                new_lines = [l for l in lines if keyword not in l]
                with open(PREF_FILE, 'w', encoding='utf-8') as f:
                    f.writelines(new_lines)
            except:
                pass
        # Also remove from dict
        for k in list(memory_facts.keys()):
            if keyword in k:
                del memory_facts[k]
        save_memory(conversation, memory_facts)
        return f"已清除包含'{keyword}'的偏好"
    else:
        if os.path.exists(PREF_FILE):
            os.remove(PREF_FILE)
        memory_facts.clear()
        save_memory(conversation, memory_facts)
        return "已清除全部偏好"


def tool_edge_browse(url, action="read"):
    """Use Edge browser (via puppeteer) to open a web page and get content."""
    EDGE_SCRIPT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "edge-browse.js")
    if not os.path.exists(EDGE_SCRIPT):
        return "[错误] 找不到 edge-browse.js"

    cmd = f'node "{EDGE_SCRIPT}" "{url}"'
    try:
        result = subprocess.run(
            cmd, shell=True, capture_output=True, text=True, timeout=45,
            cwd=os.path.dirname(EDGE_SCRIPT),
            encoding='utf-8', errors='replace',
        )

        if result.returncode != 0:
            return f"[Edge 错误] {result.stderr[:300]}"

        data = json.loads(result.stdout)

        output = f"📄 {data['title']}\n"
        output += f"🔗 {data['url']}\n"
        if data['loggedIn']:
            output += f"🔑 已登录 (Cookie: {data['cookieCount']}个)\n"
        output += "\n"

        if data['text']:
            text = data['text']
            if len(text) > 2000:
                text = text[:2000] + "\n...(共" + str(len(data['text'])) + "字，已截断。如需更多请用 web_fetch 获取完整页面)"
            output += text
        else:
            output += "(页面内容为空，可能需动态加载)"

        return output

    except subprocess.TimeoutExpired:
        return "[错误] Edge 浏览超时 (45秒)"
    except json.JSONDecodeError as e:
        return f"[错误] 解析返回数据失败: {result.stdout[:300]}"
def tool_get_time():
    """Return current date/time."""
    now = time.localtime()
    weekdays = ['一', '二', '三', '四', '五', '六', '日']
    return (
        f"现在是 {now.tm_year}年{now.tm_mon}月{now.tm_mday}日 "
        f"星期{weekdays[now.tm_wday]} "
        f"{now.tm_hour:02d}:{now.tm_min:02d}:{now.tm_sec:02d}"
    )


GITHUB_TOKEN = "你的GitHub-Token"

def tool_github_search(query, stype="repositories"):
    """Search GitHub repos or code."""
    try:
        url = f"https://api.github.com/search/{stype}?q={urllib.parse.quote(query)}&sort=stars&per_page=8"
        req = urllib.request.Request(url, headers={
            "User-Agent": "wechat-bot", "Authorization": f"Bearer {GITHUB_TOKEN}"
        })
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode())
        items = data.get("items", [])
        if not items:
            return f"未找到与 '{query}' 相关的 GitHub {stype}"
        lines = []
        for item in items[:8]:
            if stype == "repositories":
                lines.append(f"[{item['full_name']}]({item['html_url']}) ⭐{item['stargazers_count']}")
                if item.get("description"):
                    lines.append(f"  {item['description'][:150]}")
            else:
                lines.append(f"{item['repository']['full_name']}: {item['path']}")
            lines.append("")
        return "\n".join(lines)
    except Exception as e:
        return f"[GitHub搜索失败] {str(e)[:100]}"


def tool_web_search(query):
    """Search the web using Chrome browser + Bing, with HTTP fallback."""
    # Try Chrome CDP first
    CHROME_SCRIPT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "chrome-browse.js")
    if os.path.exists(CHROME_SCRIPT):
        search_url = "https://www.bing.com/search?q=" + urllib.parse.quote(query) + "&cc=cn"
        cmd = f'node "{CHROME_SCRIPT}" "{search_url}"'
        try:
            result = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=45,
                cwd=os.path.dirname(CHROME_SCRIPT), encoding='utf-8', errors='replace')
            if result.returncode == 0 and result.stdout:
                data = json.loads(result.stdout)
                text = data.get('text', '')
                if text and len(text) > 50:
                    return f"📄 {data.get('title','')}\n{text[:4000]}"
        except:
            pass

    # Fallback: HTTP request to Bing
    try:
        url = "https://www.bing.com/search?q=" + urllib.parse.quote(query) + "&cc=cn&setlang=zh-cn"
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        with urllib.request.urlopen(req, timeout=15) as resp:
            html = resp.read().decode('utf-8', errors='replace')
        # Extract text from Bing results
        results = re.findall(r'<li class="b_algo"[^>]*>.*?<h2[^>]*><a[^>]*>(.*?)</a>.*?<p[^>]*>(.*?)</p>', html, re.DOTALL)
        if results:
            lines = []
            for i, (title, snippet) in enumerate(results[:8]):
                title = re.sub(r'<[^>]+>', '', title).strip()
                snippet = re.sub(r'<[^>]+>', '', snippet).strip()
                lines.append(f"{i+1}. {title}\n   {snippet[:200]}")
            return "搜索 " + query + ":\n\n" + "\n\n".join(lines)
        return tool_web_search_lite(query)
    except Exception as e:
        try:
            return tool_web_search_lite(query)
        except:
            return f"[搜索失败] {str(e)[:100]}"


def tool_web_search_lite(query):
    """Fallback: scrape DuckDuckGo Lite HTML for search results."""
    url = "https://lite.duckduckgo.com/lite/?q=" + urllib.parse.quote(query)
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=15) as resp:
        html = resp.read().decode("utf-8", errors="replace")

    # Extract result rows: <a rel="nofollow" href="...">title</a><br><span>snippet</span>
    results = []
    # Pattern: link with title, followed by snippet in a span
    links = re.findall(r'<a[^>]*href="([^"]+)"[^>]*>([^<]+)</a>', html)
    snippets = re.findall(r'<span class="snippet">([^<]+)</span>', html)

    for i, (href, title) in enumerate(links):
        if 'duckduckgo.com' in href or not title.strip():
            continue
        snippet = snippets[i] if i < len(snippets) else ""
        clean_title = html_mod.unescape(title.strip())
        clean_snippet = html_mod.unescape(snippet.strip())
        results.append(f"[{clean_title}]({href})\n  {clean_snippet}")

    if not results:
        return f"未找到与 '{query}' 相关的结果。"

    return "\n\n".join(results[:10])


def tool_web_fetch(url):
    """Fetch and extract text from a webpage."""
    if not url.startswith(("http://", "https://")):
        url = "https://" + url

    try:
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (compatible; cc-connect-bot/1.0)"
        })
        with urllib.request.urlopen(req, timeout=20) as resp:
            # Check content type
            content_type = resp.headers.get("Content-Type", "")
            if "text" not in content_type and "html" not in content_type:
                return f"[跳过] 非文本内容类型: {content_type}"

            html = resp.read().decode("utf-8", errors="replace")

        # Strip tags, scripts, styles
        for tag in ['script', 'style', 'nav', 'footer', 'header']:
            html = re.sub(rf'<{tag}[^>]*>.*?</{tag}>', '', html, flags=re.DOTALL | re.IGNORECASE)

        text = re.sub(r'<[^>]+>', ' ', html)
        text = html_mod.unescape(text)
        text = re.sub(r'\s+', ' ', text).strip()

        if len(text) > 5000:
            text = text[:5000] + "\n...(内容已截断)"

        return text if text.strip() else "(无法提取文本内容)"

    except urllib.error.HTTPError as e:
        return f"[HTTP {e.code}] 无法访问此页面"
    except urllib.error.URLError as e:
        return f"[连接失败] {str(e.reason)}"
    except Exception as e:
        return f"[抓取失败] {str(e)}"


def tool_ask_model(model, question, system_prompt=None):
    """Query a specific AI model."""
    if model not in MODELS:
        available = ", ".join(MODELS.keys())
        return f"[错误] 未知模型 '{model}'。可用: {available}"

    config = MODELS[model]
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": question})

    try:
        body = json.dumps({
            "model": model,
            "messages": messages,
            "max_tokens": 8192,
            "temperature": 0.7,
        }).encode("utf-8")

        req = urllib.request.Request(
            config["base_url"],
            data=body,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {config['api_key']}",
            },
        )

        with urllib.request.urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            reply = _convert_emoji(data["choices"][0]["message"]["content"])
            return f"[{config['name']}] 回答:\n{reply}"

    except Exception as e:
        return f"[{config['name']}] 调用失败: {str(e)}"


def tool_analyze_video(path, frame_count=None):
    """Extract keyframes from video and analyze with DeepSeek vision."""
    if not os.path.exists(path):
        return f"[错误] 视频文件不存在: {path}"

    frame_count = min(max(frame_count or 5, 2), 10)  # 2-10 frames

    # Find ffmpeg
    ffmpeg = _find_ffmpeg()
    if not ffmpeg:
        return "[错误] 未找到 ffmpeg。请安装: winget install ffmpeg"

    # Create temp directory for frames
    tmpdir = os.path.join(os.environ.get("TEMP", "C:\\Windows\\Temp"), "video_frames")
    os.makedirs(tmpdir, exist_ok=True)

    try:
        # Get video duration
        result = subprocess.run(
            f'"{ffmpeg}" -i "{path}" 2>&1',
            shell=True, capture_output=True, text=True, timeout=30,
            errors='replace'
        )
        stderr = result.stderr
        duration_match = re.search(r'Duration:\s*(\d+):(\d+):(\d+)', stderr)
        if duration_match:
            h, m, s = map(int, duration_match.groups())
            duration_secs = h * 3600 + m * 60 + s
        else:
            duration_secs = 60  # Assume 60s if can't detect

        # Extract frames at intervals
        frames = []
        step = max(1, duration_secs // (frame_count + 1))
        for i in range(1, frame_count + 1):
            t = min(step * i, duration_secs - 1)
            frame_path = os.path.join(tmpdir, f"frame_{i:02d}.jpg")
            subprocess.run(
                f'"{ffmpeg}" -ss {t} -i "{path}" -vframes 1 -q:v 2 "{frame_path}" -y',
                shell=True, capture_output=True, timeout=30,
            )
            if os.path.exists(frame_path):
                # Convert to base64 data URL
                import base64
                with open(frame_path, 'rb') as f:
                    img_data = base64.b64encode(f.read()).decode()
                frames.append({
                    "time": t,
                    "data_url": f"data:image/jpeg;base64,{img_data}",
                })
                os.remove(frame_path)

        if not frames:
            return "[错误] 无法从视频中提取帧"

        # Analyze each frame with DeepSeek vision
        analysis_parts = [f"视频文件: {os.path.basename(path)} (时长: {duration_secs}秒)\n"]
        analysis_parts.append(f"已提取 {len(frames)} 个关键帧进行分析:\n")

        for i, frame in enumerate(frames):
            mm, ss = divmod(frame["time"], 60)
            analysis_parts.append(f"\n--- 第{i+1}帧 ({int(mm)}:{int(ss):02d}) ---")

            try:
                body = json.dumps({
                    "model": "deepseek-chat",
                    "messages": [{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": f"请用一句话描述这个视频帧的内容（这是第{i+1}/{len(frames)}帧）。描述画面中的物体、人物、场景。"},
                            {"type": "image_url", "image_url": {"url": frame["data_url"]}},
                        ]
                    }],
                    "max_tokens": 200,
                }).encode("utf-8")

                req = urllib.request.Request(
                    DEEPSEEK_URL,
                    data=body,
                    headers={
                        "Content-Type": "application/json",
                        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                    },
                )
                with urllib.request.urlopen(req, timeout=60) as resp:
                    data = json.loads(resp.read().decode("utf-8"))
                    desc = data["choices"][0]["message"]["content"]
                    analysis_parts.append(desc)

            except Exception as e:
                analysis_parts.append(f"(分析失败: {e})")

        # Final summary: send all frame descriptions to DeepSeek
        frame_descriptions = "\n".join(
            p for p in analysis_parts[2:] if not p.startswith("(")
        )
        try:
            body = json.dumps({
                "model": "deepseek-chat",
                "messages": [{
                    "role": "user",
                    "content": f"以下是一个视频各帧的描述，请用中文总结这个视频的内容（100字以内）:\n\n{frame_descriptions}",
                }],
                "max_tokens": 300,
            }).encode("utf-8")

            req = urllib.request.Request(
                DEEPSEEK_URL,
                data=body,
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                },
            )
            with urllib.request.urlopen(req, timeout=60) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                summary = data["choices"][0]["message"]["content"]
                analysis_parts.append(f"\n## 视频总结\n{summary}")

        except Exception as e:
            analysis_parts.append(f"\n(总结失败: {e})")

        return "\n".join(analysis_parts)

    except Exception as e:
        return f"[视频分析失败] {str(e)}"


def _find_ffmpeg():
    """Find ffmpeg executable."""
    import shutil as _shutil
    # Check PATH first
    found = _shutil.which("ffmpeg")
    if found:
        return found
    # Check imageio-ffmpeg (Python package)
    try:
        import imageio_ffmpeg
        exe = imageio_ffmpeg.get_ffmpeg_exe()
        if os.path.exists(exe):
            return exe
    except ImportError:
        pass
    # Check common install locations on Windows
    for loc in [
        r"C:\ffmpeg\bin\ffmpeg.exe",
        r"C:\Program Files\ffmpeg\bin\ffmpeg.exe",
        os.path.expandvars(r"%LOCALAPPDATA%\Microsoft\WinGet\Packages\Gyan.FFmpeg_*\ffmpeg.exe"),
    ]:
        matches = glob.glob(loc)
        if matches:
            return matches[0]
    return None


def tool_read_docx(path):
    """Read a Word document, extracting text and describing embedded images."""
    if not os.path.exists(path):
        return f"[错误] 文件不存在: {path}"
    if not path.lower().endswith('.docx'):
        return "[错误] 仅支持 .docx 格式"

    try:
        import docx
        import base64
        import tempfile

        doc = docx.Document(path)
        result_parts = [f"Word 文档: {os.path.basename(path)}\n"]

        # Extract text from paragraphs
        text_lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_lines.append(para.text.strip())
        if text_lines:
            result_parts.append("=== 文字内容 ===\n" + "\n".join(text_lines))

        # Extract tables
        if doc.tables:
            result_parts.append(f"\n=== 表格 ({len(doc.tables)}个) ===")
            for i, table in enumerate(doc.tables):
                result_parts.append(f"\n--- 表格{i+1} ---")
                for j, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    result_parts.append(" | ".join(cells))
                    if j > 50:
                        result_parts.append("...(表格过大，已截断)")
                        break

        # Extract and describe images
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
                if image_count > 10:
                    break
                try:
                    img_data = rel.target_part.blob
                    data_url = f"data:image/png;base64,{base64.b64encode(img_data).decode()}"

                    body = json.dumps({
                        "model": "deepseek-chat",
                        "messages": [{
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "请用一句话描述这张图片的内容"},
                                {"type": "image_url", "image_url": {"url": data_url}},
                            ]
                        }],
                        "max_tokens": 100,
                    }).encode("utf-8")

                    req = urllib.request.Request(
                        DEEPSEEK_URL, data=body,
                        headers={"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
                    )
                    with urllib.request.urlopen(req, timeout=30) as resp:
                        data = json.loads(resp.read().decode("utf-8"))
                        desc = data["choices"][0]["message"]["content"]
                        result_parts.append(f"\n[图片{image_count}]: {desc}")
                except Exception as e:
                    result_parts.append(f"\n[图片{image_count}]: (识别失败: {e})")

        if image_count == 0:
            result_parts.append("\n(文档中无嵌入图片)")
        elif image_count > 10:
            result_parts.append(f"\n(共{image_count}张图片，仅分析了前10张)")

        return "\n".join(result_parts)

    except ImportError:
        return "[错误] python-docx 未安装。运行: pip install python-docx"
    except Exception as e:
        return f"[读取失败] {str(e)}"


    except Exception as e:
        return f"[错误] {str(e)[:200]}"


def tool_bash(command):
    """Execute a shell command."""
    if SAFE_MODE:
        # Safety: block dangerous commands
        dangerous = ['rm -rf /', 'format', 'del /f /s', 'shutdown', 'restart', 'logoff']
        cmd_lower = command.lower()
        for d in dangerous:
            if d in cmd_lower:
                return f"[拒绝] 命令包含危险操作: {d}"

    try:
        result = subprocess.run(
            command,
            shell=True,
            capture_output=True,
            text=True,
            timeout=120,
            cwd="C:\\Users\\big",
            encoding='utf-8',
            errors='replace',
        )
        output = result.stdout
        if result.stderr:
            output += "\n[stderr]\n" + result.stderr
        if not output.strip():
            output = f"(命令执行完毕，退出码: {result.returncode})"
        # Truncate if too long
        if len(output) > 4000:
            output = output[:1500] + f"\n...(共{len(output)}字已截断，需要更多请用更精确的命令)"
        return output
    except subprocess.TimeoutExpired:
        return "[错误] 命令超时 (120秒)"
    except Exception as e:
        return f"[错误] {str(e)}"


def tool_list_dir(path):
    """List directory contents."""
    if not is_path_allowed(path):
        return f"[拒绝] 无权访问: {path}"
    try:
        items = os.listdir(path)
        if not items:
            return "(空目录)"
        lines = []
        for item in sorted(items):
            full = os.path.join(path, item)
            try:
                if os.path.isdir(full):
                    lines.append(f"[DIR]  {item}/")
                else:
                    size = os.path.getsize(full)
                    if size < 1024:
                        size_str = f"{size}B"
                    elif size < 1024 * 1024:
                        size_str = f"{size / 1024:.1f}KB"
                    else:
                        size_str = f"{size / (1024 * 1024):.1f}MB"
                    lines.append(f"[FILE] {item} ({size_str})")
            except OSError:
                lines.append(f"[???]  {item}")
        return "\n".join(lines[:200])
    except PermissionError:
        return f"[错误] 无权限访问: {path}"
    except FileNotFoundError:
        return f"[错误] 目录不存在: {path}"
    except Exception as e:
        return f"[错误] {str(e)}"


def tool_grep(pattern, path):
    """Search for pattern in files."""
    if not is_path_allowed(path):
        return f"[拒绝] 无权访问: {path}"
    try:
        results = []
        if os.path.isdir(path):
            files = glob.glob(os.path.join(path, "**", "*"), recursive=True)
            # Filter to text files only
            text_exts = {'.txt', '.py', '.js', '.ts', '.html', '.css', '.json', '.xml',
                        '.md', '.yml', '.yaml', '.toml', '.ini', '.cfg', '.log', '.csv',
                        '.bat', '.cmd', '.ps1', '.sh', '.java', '.c', '.cpp', '.h', '.rs',
                        '.go', '.rb', '.php', '.sql'}
            files = [f for f in files if os.path.splitext(f)[1].lower() in text_exts][:100]
            if not files:
                return f"[提示] 未找到文本文件在: {path}"
        else:
            files = [path]

        for filepath in files:
            try:
                with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                    for i, line in enumerate(f, 1):
                        if pattern.lower() in line.lower():
                            results.append(f"{filepath}:{i}: {line.strip()}")
                            if len(results) > 50:
                                break
            except Exception:
                continue
            if len(results) > 50:
                break

        if not results:
            return f"未找到匹配 '{pattern}' 的内容"
        return "\n".join(results[:50]) + ("\n... (仅显示前50条)" if len(results) > 50 else "")
    except Exception as e:
        return f"[错误] {str(e)}"


# Tool dispatcher
TOOL_FUNCTIONS = {
    "remember_preference": lambda args: tool_remember_preference(args.get("rule", "")),
    "forget_preferences": lambda args: tool_forget_preferences(args.get("keyword")),
    "web_search": lambda args: tool_web_search(args.get("query", "")),
    "github_search": lambda args: tool_github_search(args.get("query", ""), args.get("type", "repositories")),
    "get_time": lambda args: tool_get_time(),
    "ask_model": lambda args: tool_ask_model(args.get("model", ""), args.get("question", ""), None),
    "send_to_wechat": lambda args: tool_send_to_wechat(args.get("path", ""), args.get("type", "file")),
    "bash": lambda args: tool_bash(args.get("command", "")),
    "read": lambda args: tool_read(args.get("path", ""), args.get("limit", 200)),
    "write": lambda args: tool_write(args.get("path", ""), args.get("content", "")),
    "list_dir": lambda args: tool_list_dir(args.get("path", "")),
    "grep": lambda args: tool_grep(args.get("pattern", ""), args.get("path", "")),
}


# ============================================================
# DeepSeek API call with tool support
# ============================================================

def deepseek_chat_with_tools(messages, model=None):
    """Call DeepSeek API with tool definitions, return full response."""
    model = model or DEFAULT_MODEL
    api_key = MODELS.get(model, MODELS[DEFAULT_MODEL])["api_key"]
    base_url = MODELS.get(model, MODELS[DEFAULT_MODEL])["base_url"]

    body = json.dumps({
        "model": model,
        "messages": messages,
        "tools": TOOLS,
        "max_tokens": 4096,  # 首次调用用小的（闲聊够用）
        "temperature": 0.7,
    }).encode("utf-8")

    req = urllib.request.Request(
        base_url,
        data=body,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        },
    )

    with urllib.request.urlopen(req, timeout=120) as resp:
        return json.loads(resp.read().decode("utf-8"))


# ============================================================
# Stream-json output helpers
# ============================================================

def send_json(obj):
    """Send a JSON object to stdout."""
    sys.stdout.write(json.dumps(obj, ensure_ascii=False) + "\n")
    sys.stdout.flush()


def send_tool_event(tool_name, tool_input, tool_result, session_id, msg_id):
    """Send tool call and result as stream-json events."""
    tool_id = "toolu_" + str(uuid.uuid4())[:12]

    # Tool call
    send_json({
        "type": "assistant",
        "message": {
            "id": msg_id,
            "type": "message",
            "role": "assistant",
            "model": DEFAULT_MODEL,
            "content": [{
                "type": "tool_use",
                "id": tool_id,
                "name": tool_name,
                "input": tool_input,
            }],
            "stop_reason": "tool_use",
        },
        "session_id": session_id,
        "uuid": str(uuid.uuid4()),
    })

    # Tool result (as user message)
    send_json({
        "type": "user",
        "message": {
            "role": "user",
            "content": [{
                "type": "tool_result",
                "tool_use_id": tool_id,
                "content": tool_result,
            }],
        },
        "session_id": session_id,
        "uuid": str(uuid.uuid4()),
    })


# ============================================================
# Main message handler with agent loop
# ============================================================

def handle_user_message(text, images=None):
    """Handle a user message with agent loop (tool calling)."""
    global conversation, memory_facts

    if images is None:
        images = []

    # Handle /reset command
    if text and text.strip().startswith("/reset"):
        conversation.clear()
        memory_facts.clear()
        save_memory(conversation, memory_facts)
        if os.path.exists(PREF_FILE):
            os.remove(PREF_FILE)
        reply = "✅ 已重置：对话历史、偏好设置全部清除。"
        conversation.append({"role": "user", "content": text})
        conversation.append({"role": "assistant", "content": reply})
        send_json({"type": "assistant", "message": {"role": "assistant", "content": [{"type": "text", "text": reply}]}})
        send_json({"type": "result", "subtype": "success", "is_error": False, "result": reply})
        return

    session_id = "ses_" + str(uuid.uuid4())[:12]
    msg_id = "msg_" + str(uuid.uuid4())[:12]

    # Build user content
    user_content = []
    if text:
        user_content.append({"type": "text", "text": text})
    for i, img in enumerate(images):
        try:
            desc = _describe_image(img)
            user_content.append({"type": "text", "text": f"[图片{i+1}] {desc}"})
        except Exception as e:
            user_content.append({"type": "text", "text": f"[图片{i+1}: 未能识别]"})

    # Smart history selection based on context
    conv_len = len(conversation)
    if conv_len < SHORT_HISTORY:
        history_limit = SHORT_HISTORY
    elif conv_len < NORMAL_HISTORY * 2:
        history_limit = NORMAL_HISTORY
    else:
        history_limit = LONG_HISTORY

    recent_messages = conversation[-history_limit:] if len(conversation) > history_limit else conversation
    api_messages = [{"role": "system", "content": get_system_prompt()}]
    api_messages.extend(recent_messages)
    api_messages.append({"role": "user", "content": user_content if images else text})

    max_tool_rounds = 100
    tool_count = 0

    try:
        while tool_count < max_tool_rounds:
            response = deepseek_chat_with_tools(api_messages)
            choice = response["choices"][0]
            message = choice["message"]
            finish_reason = choice.get("finish_reason", "stop")

            # Check if the model wants to call a tool
            if finish_reason == "tool_calls" or message.get("tool_calls"):
                tool_calls = message.get("tool_calls", [])

                # Add assistant message (with tool calls) to history
                api_messages.append({
                    "role": "assistant",
                    "content": message.get("content"),
                    "tool_calls": tool_calls,
                })

                for tc in tool_calls:
                    fn_name = tc["function"]["name"]
                    fn_args = json.loads(tc["function"]["arguments"])
                    tool_count += 1

                    # Execute tool
                    fn = TOOL_FUNCTIONS.get(fn_name)
                    if fn:
                        result = fn(fn_args)
                    elif _mcp_available:
                        # Try MCP tools
                        result = call_mcp_tool(fn_name, fn_args)
                    else:
                        result = f"[错误] 未知工具: {fn_name}"

                    # Send tool events to cc-connect
                    send_tool_event(fn_name, fn_args, result, session_id, msg_id)

                    # Add tool result to API messages
                    api_messages.append({
                        "role": "tool",
                        "tool_call_id": tc["id"],
                        "content": result,
                    })
            else:
                # Final text response
                reply = message.get("content", "") or "(空响应)"
                # Auto-convert WeChat bracket emoji codes to Unicode emoji
                reply = _convert_emoji(reply)

                # Update conversation (keep ALL history in memory + disk)
                # For image messages, save text + 📷 marker (image URLs are temporary)
                if images:
                    conversation.append({"role": "user", "content": (text or "请描述图片") + " 📷[图片]"})
                else:
                    conversation.append({"role": "user", "content": text})
                conversation.append({"role": "assistant", "content": reply})

                # Auto-save to disk
                save_memory(conversation, memory_facts)

                # Send thinking
                send_json({
                    "type": "assistant",
                    "message": {
                        "id": msg_id,
                        "type": "message",
                        "role": "assistant",
                        "model": DEFAULT_MODEL,
                        "content": [{"type": "thinking", "thinking": "", "signature": ""}],
                        "stop_reason": None,
                        "usage": {"input_tokens": 0, "output_tokens": 0},
                    },
                    "session_id": session_id,
                    "uuid": str(uuid.uuid4()),
                })

                # Send text response
                send_json({
                    "type": "assistant",
                    "message": {
                        "id": msg_id,
                        "type": "message",
                        "role": "assistant",
                        "model": DEFAULT_MODEL,
                        "content": [{"type": "text", "text": reply}],
                        "stop_reason": "end_turn",
                        "usage": {"input_tokens": 0, "output_tokens": 0},
                    },
                    "session_id": session_id,
                    "uuid": str(uuid.uuid4()),
                })

                # Send result
                send_json({
                    "type": "result",
                    "subtype": "success",
                    "is_error": False,
                    "duration_ms": 1000,
                    "num_turns": 1 + tool_count,
                    "result": reply,
                    "stop_reason": "end_turn",
                    "session_id": session_id,
                    "total_cost_usd": 0,
                    "usage": {"input_tokens": 0, "output_tokens": 0},
                    "permission_denials": [],
                })
                return  # Done

        # Exceeded max tool rounds
        error_text = "[提示] 工具调用次数过多，已停止。请简化你的请求。"
        send_json({
            "type": "result",
            "subtype": "error",
            "is_error": True,
            "result": error_text,
            "session_id": session_id,
        })

    except urllib.error.HTTPError as e:
        error_body = e.read().decode("utf-8", errors="replace")
        error_text = f"[API Error] HTTP {e.code}: {error_body[:300]}"
        send_json({
            "type": "result",
            "subtype": "error",
            "is_error": True,
            "result": error_text,
            "session_id": session_id,
        })
    except Exception as e:
        error_text = f"[Error] {str(e)}"
        send_json({
            "type": "result",
            "subtype": "error",
            "is_error": True,
            "result": error_text,
            "session_id": session_id,
        })


# ============================================================
# Input processing
# ============================================================

def _convert_emoji(text):
    """Convert WeChat bracket-format emoji codes to Unicode emoji."""
    EMOJI_MAP = {
        '[旺柴]': '🐕', '[偷笑]': '🤭', '[笑哭]': '😂', '[捂脸]': '🤦',
        '[耶]': '✌️', '[加油]': '💪', '[好的]': '👌', '[OK]': '👌',
        '[赞]': '👍', '[爱心]': '❤️', '[心]': '❤️', '[玫瑰]': '🌹',
        '[握手]': '🤝', '[抱拳]': '🙏', '[合十]': '🙏', '[谢谢]': '🙏',
        '[强]': '💪', '[弱]': '👎', '[胜利]': '✌️', '[拳头]': '👊',
        '[鼓掌]': '👏', '[派对]': '🎉', '[烟花]': '🎆', '[爆竹]': '🧨',
        '[红包]': '🧧', '[礼物]': '🎁', '[蜡烛]': '🕯️', '[蛋糕]': '🎂',
        '[咖啡]': '☕', '[啤酒]': '🍺', '[干杯]': '🥂', '[饭]': '🍚',
        '[猪头]': '🐷', '[哈欠]': '🥱', '[委屈]': '🥺', '[哭]': '😢',
        '[笑哭]': '😂', '[憨笑]': '😁', '[微笑]': '🙂', '[色]': '😍',
        '[发呆]': '😳', '[得意]': '😏', '[呲牙]': '😬', '[惊讶]': '😲',
        '[难过]': '😔', '[酷]': '😎', '[冷汗]': '😰', '[抓狂]': '😫',
        '[吐]': '🤮', '[睡]': '😴', '[调皮]': '😜', '[白眼]': '🙄',
        '[傲慢]': '😤', '[困]': '😪', '[惊恐]': '😱', '[流汗]': '😅',
        '[憨笑]': '😁', '[悠闲]': '😌', '[奋斗]': '💪', '[咒骂]': '🤬',
        '[疑问]': '❓', '[嘘]': '🤫', '[晕]': '😵', '[敲打]': '🤛',
        '[再见]': '👋', '[擦汗]': '😅', '[抠鼻]': '🤏', '[鼓掌]': '👏',
        '[坏笑]': '😏', '[左哼哼]': '😤', '[右哼哼]': '😤', '[鄙视]': '😒',
        '[阴险]': '😈', '[亲亲]': '😘', '[可怜]': '🥺', '[笑脸]': '😊',
        '[生病]': '🤒', '[口红]': '💄', '[太阳]': '☀️', '[月亮]': '🌙',
        '[下雨]': '🌧️', '[雪]': '❄️', '[风]': '💨', '[闪电]': '⚡',
        '[庆祝]': '🎉', '[福]': '🧧', '[发]': '💰', '[红包]': '🧧',
        '[炸弹]': '💣', '[刀]': '🔪', '[西瓜]': '🍉', '[OK]': '👌',
    }
    for code, emoji in EMOJI_MAP.items():
        text = text.replace(code, emoji)
    return text


def debug_tag(msg, tag):
    """Write a debug entry to debug.jsonl."""
    entry = dict(tag=tag, time=time.strftime("%H:%M:%S"), data=msg)
    try:
        with open(DEBUG_LOG, 'a', encoding='utf-8') as f:
            f.write(json.dumps(entry, ensure_ascii=False) + '\n')
    except:
        pass


def process_input(line):
    """Process one line of stream-json input."""
    line = line.strip()
    if not line:
        return

    # Debug: dump raw input
    debug_tag(line[:200] if len(line) > 10000 else line, "raw_input")

    try:
        msg = json.loads(line)
    except json.JSONDecodeError as e:
        sys.stderr.write(f"JSON parse error: {e}\n")
        sys.stderr.flush()
        if '"content"' in line or '"text"' in line or '"source"' in line:
            try:
                fixed = line.replace('\r', '\\r').replace('\n', '\\n')
                msg = json.loads(fixed)
            except:
                return
        else:
            return

    msg_type = msg.get("type", "")
    message = msg.get("message", {})
    content = message.get("content", [])

    # Handle image messages (may not have type="user")
    if isinstance(content, list):
        for part in content:
            if isinstance(part, dict) and "source" in part:
                src = part.get("source", {})
                img_data = src.get("data", "")
                if img_data and src.get("media_type", "").startswith("image"):
                    try:
                        import base64
                        mime = src.get("media_type", "image/jpeg")
                        ext = mime.split("/")[-1] if "/" in mime else "jpg"
                        img_path = os.path.join(ATTACH_DIR, f"img_stream_{int(time.time()*1000)}.{ext}")
                        with open(img_path, "wb") as f:
                            f.write(base64.b64decode(img_data))
                        desc = _describe_image_qwen(img_path)
                        handle_user_message(f"[微信图片内容: {desc}]")
                        return
                    except Exception as e:
                        sys.stderr.write(f"[IMG] Decode failed: {e}\n")
                        sys.stderr.flush()

    if msg_type != "user":
        sys.stderr.write(f"DEBUG msg_type={msg_type} keys={list(msg.keys())[:8]}\n")
        sys.stderr.flush()

    if msg_type == "user":

        # Full dump for debugging file/image messages
        has_file = any(isinstance(p, dict) and p.get("type") not in ("text",) for p in (content if isinstance(content, list) else []))
        if has_file:
            sys.stderr.write(f"DEBUG FILE_MSG: {json.dumps(msg, ensure_ascii=False)[:2000]}\n")
            sys.stderr.flush()

        text_parts = []
        images = []

        # Check for image attachments from cc-connect's meta
        # If user sent an image, wait briefly then scan attachments dir
        has_images_meta = msg.get("has_images") or msg.get("message", {}).get("has_images", False) or has_file or (not text_parts and not content)
        if has_images_meta:
            # Wait for cc-connect to finish downloading the image
            time.sleep(0.5)
            # Force scan attachments now (watcher runs every 3s, too slow)
            _force_scan_attachments()

        if isinstance(content, str):
            content = [{"type": "text", "text": content}]

        if not isinstance(content, list):
            handle_user_message(str(content), [])
            return

        for part in content:
            if isinstance(part, dict):
                ptype = part.get("type", "")
                if ptype == "text":
                    text_parts.append(part.get("text", ""))
                elif ptype in ("image", "image_url"):
                    img_url = _extract_image_url(part)
                    if img_url:
                        images.append(img_url)
                elif ptype == "file":
                    # File attachment - check if it's an image/video
                    file_path = part.get("path", "") or part.get("source", {}).get("path", "")
                    if file_path and os.path.exists(file_path):
                        ext = os.path.splitext(file_path)[1].lower()
                        if ext in ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'):
                            img_url = _file_to_data_url(file_path)
                            if img_url:
                                images.append(img_url)
                                text_parts.append(f"[收到图片: {os.path.basename(file_path)}]")
                        elif ext in ('.mp4', '.avi', '.mov', '.mkv', '.webm'):
                            text_parts.append(f"[收到视频: {os.path.basename(file_path)}] 请用 analyze_video 工具分析")
                        else:
                            text_parts.append(f"[收到文件: {os.path.basename(file_path)}]")
                    else:
                        text_parts.append("[收到文件，但无法访问]")
                elif ptype == "tool_result":
                    pass
                else:
                    sys.stderr.write(f"UNKNOWN content type: {json.dumps(part, ensure_ascii=False)[:500]}\n")
                    sys.stderr.flush()

        user_text = " ".join(text_parts).strip()

        # Inject any pending image descriptions from watcher thread
        img_text = get_pending_image_text()
        if img_text and user_text:
            user_text = user_text + "\n\n" + img_text
        elif img_text and not user_text:
            user_text = img_text
        # Also check for images in the cc-connect message
        if user_text or images:
            handle_user_message(user_text or "请描述这张图片", images)


def _extract_image_url(part):
    """Extract image URL from various message formats."""
    # Format 1: {"type": "image", "source": {"type": "url", "url": "..."}}
    source = part.get("source", {})
    if isinstance(source, dict):
        if source.get("type") == "url" and source.get("url"):
            return source["url"]
        if source.get("type") == "base64" and source.get("data"):
            return f"data:{source.get('media_type', 'image/jpeg')};base64,{source['data']}"
        if source.get("type") == "file" and source.get("path"):
            return _file_to_data_url(source["path"])

    # Format 2: {"type": "image_url", "image_url": {"url": "..."}}
    image_url = part.get("image_url", {})
    if isinstance(image_url, dict) and image_url.get("url"):
        return image_url["url"]

    # Format 3: {"type": "image", "url": "..."}
    if part.get("url"):
        return part["url"]

    # Format 4: cc-connect sometimes saves image and sends path
    # Look for path in various places
    path = part.get("path", "")
    if path:
        return _resolve_image_path(path)

    return None


# Track files seen by force-scan so we don't re-inject
_force_scan_seen = set()
_force_scan_lock = threading.Lock()


def _force_scan_attachments():
    """Immediately scan attachments dir for new images, bypassing watcher delay."""
    global _force_scan_seen
    if not os.path.isdir(ATTACH_DIR):
        return
    now = time.time()
    with _force_scan_lock:
        for fname in sorted(os.listdir(ATTACH_DIR), key=lambda x: os.path.getmtime(os.path.join(ATTACH_DIR, x)) if os.path.exists(os.path.join(ATTACH_DIR, x)) else 0):
            fpath = os.path.join(ATTACH_DIR, fname)
            if fpath in _force_scan_seen:
                continue
            _force_scan_seen.add(fpath)
            try:
                st = os.stat(fpath)
            except OSError:
                continue
            if st.st_size < 5000 or now - st.st_mtime < 0.5:
                continue
            ext = os.path.splitext(fname)[1].lower()
            if ext not in ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'):
                continue
            # Found a new image - analyze with Qwen-VL then add to pending
            try:
                description = _describe_image_qwen(fpath)
            except Exception:
                description = "(图片识别失败)"
            with _pending_image_lock:
                _pending_image_descriptions.append(
                    f"[来自微信的图片: {fname}]\n{description}"
                )


def _file_to_data_url(filepath):
    """Convert a local image file to base64 data URL."""
    import base64
    if not os.path.exists(filepath):
        return None
    ext = os.path.splitext(filepath)[1].lower()
    mime_map = {'.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
                '.gif': 'image/gif', '.webp': 'image/webp', '.bmp': 'image/bmp'}
    mime = mime_map.get(ext, 'image/jpeg')
    try:
        with open(filepath, 'rb') as f:
            data = base64.b64encode(f.read()).decode()
        return f"data:{mime};base64,{data}"
    except Exception:
        return None


def _resolve_image_path(path):
    """Resolve image file path - cc-connect saves to attachments dir."""
    # If it's already absolute and exists
    if os.path.isabs(path) and os.path.exists(path):
        return _file_to_data_url(path)

    # If it's a relative path starting with .cc-connect
    # The attachments are saved relative to the cc-connect working directory
    if path.startswith(".cc-connect") or "attachments" in path:
        # Try user's .cc-connect directory
        abs_path = os.path.join(os.environ.get("USERPROFILE", "C:\\Users\\big"), path.replace("./", ""))
        if os.path.exists(abs_path):
            return _file_to_data_url(abs_path)

    # Try within user's .cc-connect/attachments/
    base_attachments = os.path.join(os.environ.get("USERPROFILE", "C:\\Users\\big"), ".cc-connect", "attachments")
    basename = os.path.basename(path)
    candidate = os.path.join(base_attachments, basename)
    if os.path.exists(candidate):
        return _file_to_data_url(candidate)

    # Try wechat-bot dir
    bot_dir = os.path.join(os.environ.get("USERPROFILE", "C:\\Users\\big"), "wechat-bot")
    candidate2 = os.path.join(bot_dir, path)
    if os.path.exists(candidate2):
        return _file_to_data_url(candidate2)

    return None


def _describe_image_qwen(img_path):
    """Describe an image using Qwen-VL (阿里千问) vision API."""
    if not os.path.exists(img_path):
        # Try as data URL
        if img_path.startswith('data:'):
            image_data = img_path
        elif img_path.startswith(('http://', 'https://')):
            image_data = {"type": "image_url", "image_url": {"url": img_path}}
            return _call_qwen_vl(image_data)
        else:
            return "(图片文件不存在)"
    else:
        # Convert local file to base64
        data_url = _file_to_data_url(img_path)
        if not data_url:
            return "(无法读取图片)"

    image_data = {"type": "image_url", "image_url": {"url": data_url}}
    return _call_qwen_vl(image_data)


def _call_qwen_vl(image_data):
    """Call Qwen-VL API to analyze an image."""
    try:
        body = json.dumps({
            "model": QWEN_VL_MODEL,
            "messages": [{
                "role": "user",
                "content": [
                    image_data,
                    {"type": "text", "text": "请用一句话描述这张图片的内容，包括主要物体、场景、人物和文字（如果有）。不超过100字。"}
                ]
            }],
            "max_tokens": 300,
        }).encode("utf-8")

        req = urllib.request.Request(
            f"{QWEN_BASE_URL}/chat/completions",
            data=body,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {QWEN_API_KEY}",
            },
        )

        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return data["choices"][0]["message"]["content"]
    except urllib.error.HTTPError as e:
        error_body = e.read().decode("utf-8", errors="replace")
        sys.stderr.write(f"Qwen-VL HTTP Error {e.code}: {error_body[:200]}\n")
        sys.stderr.flush()
        return f"(视觉识别失败: HTTP {e.code})"
    except Exception as e:
        sys.stderr.write(f"Qwen-VL Error: {e}\n")
        sys.stderr.flush()
        return f"(视觉识别失败: {str(e)[:50]})"


def _describe_image(img_source):
    """Describe an image using DeepSeek V4 vision capabilities."""
    # Determine if source is URL or base64
    if img_source.startswith('data:'):
        # Already base64, use as-is
        image_content = {"type": "image_url", "image_url": {"url": img_source}}
    elif img_source.startswith(('http://', 'https://')):
        # URL
        image_content = {"type": "image_url", "image_url": {"url": img_source}}
    else:
        return "(无法识别的图片格式)"

    # Use deepseek-chat (V4 models) which support vision
    try:
        body = json.dumps({
            "model": "deepseek-chat",
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "text", "text": "请用一句话描述这张图片的内容，包括主要物体、场景、文字（如果有）。"},
                    image_content,
                ]
            }],
            "max_tokens": 200,
        }).encode("utf-8")

        req = urllib.request.Request(
            DEEPSEEK_URL, data=body,
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
            method='POST',
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return data["choices"][0]["message"]["content"]
    except Exception as e:
        sys.stderr.write(f"Image analysis failed: {e}\n")
        sys.stderr.flush()
        # Fallback: try with vision-specific model
        try:
            body = json.dumps({
                "model": "deepseek-v4-flash",
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "描述这张图片"},
                        image_content,
                    ]
                }],
                "max_tokens": 200,
            }).encode("utf-8")

            req = urllib.request.Request(
                f"https://api.deepseek.com/v1/chat/completions",
                data=body,
                headers={"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
                method='POST',
            )
            with urllib.request.urlopen(req, timeout=30) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                return data["choices"][0]["message"]["content"]
        except Exception as e2:
            return f"(图片识别失败: 请描述这张图片中的内容)"


# ============================================================
# Attachment Watcher - watches cc-connect's download folder
# ============================================================

# Global queue for images detected by watcher thread
_pending_image_descriptions = []
_pending_image_lock = threading.Lock()


def start_attachment_watcher():
    """
    Start a background thread that watches the cc-connect attachments directory.
    When new images are detected, they are described and queued for the main thread.
    """
    seen_files = set()
    if os.path.isdir(ATTACH_DIR):
        for f in os.listdir(ATTACH_DIR):
            seen_files.add(os.path.join(ATTACH_DIR, f))

    cooldown = set()

    def watcher():
        while True:
            try:
                if not os.path.isdir(ATTACH_DIR):
                    time.sleep(3)
                    continue

                now = time.time()
                for fname in os.listdir(ATTACH_DIR):
                    fpath = os.path.join(ATTACH_DIR, fname)
                    if fpath in seen_files or fpath in cooldown:
                        continue
                    seen_files.add(fpath)

                    try:
                        st = os.stat(fpath)
                    except OSError:
                        continue
                    if st.st_size < 5000 or now - st.st_mtime < 2:
                        continue
                    ext = os.path.splitext(fname)[1].lower()
                    if ext not in ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'):
                        continue

                    cooldown.add(fpath)
                    # Use Qwen-VL to analyze image
                    try:
                        description = _describe_image_qwen(fpath)
                    except Exception:
                        description = "(图片识别失败)"
                    with _pending_image_lock:
                        _pending_image_descriptions.append(
                            f"[来自微信的图片: {fname}]\n{description}"
                        )
                    sys.stderr.write(f"[IMG] Qwen-VL analyzed: {fname}\n")
                    sys.stderr.flush()
            except Exception:
                pass
            time.sleep(3)

    t = threading.Thread(target=watcher, daemon=True, name="img-watcher")
    t.start()
    sys.stderr.write(f"Attachment watcher started: {ATTACH_DIR}\n")
    sys.stderr.flush()


def get_pending_image_text():
    """Get any pending image descriptions from the watcher thread (main thread only)."""
    with _pending_image_lock:
        if _pending_image_descriptions:
            items = list(_pending_image_descriptions)
            _pending_image_descriptions.clear()
            return "\n\n".join(items)
    return None


def main():
    global conversation, memory_facts

    # Load memory from disk
    conversation, memory_facts = load_memory()

    # Write debug header
    debug_tag({"version": "1.0", "attachments": ATTACH_DIR}, "init")

    start_attachment_watcher()

    # Initialize MCP servers
    if _mcp_available and MCP_SERVERS:
        mcp_tools = init_mcp(MCP_SERVERS)
        if mcp_tools:
            for t in mcp_tools:
                TOOLS.append(t)
            sys.stderr.write(f"MCP tools added: {len(mcp_tools)} from {len(MCP_SERVERS)} servers\n")
            sys.stderr.flush()

    send_json({
        "type": "system",
        "subtype": "init",
        "cwd": os.getcwd(),
        "session_id": "ses_" + str(uuid.uuid4())[:12],
        "tools": ["Bash", "Read", "Write", "Grep", "Glob", "WebSearch", "WebFetch"],
        "mcp_servers": [],
        "model": DEFAULT_MODEL,
        "permissionMode": "bypassPermissions",
        "slash_commands": [],
        "apiKeySource": "none",
        "claude_code_version": "agent-1.0.0",
        "output_style": "default",
        "agents": ["general-purpose"],
        "skills": [],
        "plugins": [],
    })

    sys.stderr.write("DeepSeek Agent ready (tool-calling enabled)\n")
    sys.stderr.flush()

    # Main loop: process stdin lines from cc-connect
    try:
        for line in sys.stdin:
            process_input(line)
    finally:
        if _mcp_available:
            stop_all_mcp()


if __name__ == "__main__":
    main()
